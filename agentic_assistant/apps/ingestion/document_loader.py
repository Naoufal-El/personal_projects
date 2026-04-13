"""
Document Loader
Loads and extracts text from multiple file formats
"""
import json
import csv
from pathlib import Path
from typing import Dict, Tuple, Any

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class DocumentLoader:
    """
    Unified document loader for multiple file formats

    Supported formats:
    - .txt: Plain text files
    - .json: JSON documents
    - .csv: CSV files
    - .pdf: PDF documents (requires pypdf)
    - .docx: Microsoft Word documents (requires python-docx)
    """

    def __init__(self):
        self.loaders = {
            '.txt': self._load_txt,
            '.json': self._load_json,
            '.csv': self._load_csv,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
        }

    def load(self, file_path: Path) -> Dict[str, Any]:
        """
        Load document and extract text content

        Args:
            file_path: Path to document file

        Returns:
            Dict with extracted content and metadata:
            {
                'content': str,           # Extracted text
                'format': str,            # File format
                'filename': str,          # Original filename
                'size_bytes': int,        # File size
                'metadata': Dict,         # Format-specific metadata
            }

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
            Exception: If loading fails
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension not in self.loaders:
            raise ValueError(f"Unsupported file format: {extension}")

        loader = self.loaders[extension]

        try:
            content, metadata = loader(file_path)

            return {
                'content': content,
                'format': extension,
                'filename': file_path.name,
                'size_bytes': file_path.stat().st_size,
                'metadata': metadata or {},
            }

        except Exception as e:
            raise Exception(f"Failed to load {file_path.name}: {str(e)}")

    def _load_txt(self, file_path: Path) -> Tuple[str, Dict]:
        """Load plain text file"""
        try:
            with open(str(file_path), 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(str(file_path), 'r', encoding='latin-1') as f:
                content = f.read()

        metadata = {
            'encoding': 'utf-8',
            'lines': len(content.splitlines())
        }

        return content, metadata

    def _load_json(self, file_path: Path) -> Tuple[str, Dict]:
        """Load JSON file"""
        with open(str(file_path), 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert JSON to readable text
        if isinstance(data, dict):
            content = self._dict_to_text(data)
        elif isinstance(data, list):
            content = '\n\n'.join([self._dict_to_text(item) if isinstance(item, dict) else str(item) for item in data])
        else:
            content = json.dumps(data, indent=2)

        metadata = {
            'type': type(data).__name__,
            'keys': list(data.keys()) if isinstance(data, dict) else None,
            'items': len(data) if isinstance(data, list) else None
        }

        return content, metadata

    def _load_csv(self, file_path: Path) -> Tuple[str, Dict]:
        """Load CSV file"""
        with open(str(file_path), 'r', encoding='utf-8') as f:
            csv_reader = csv.DictReader(f)
            rows = list(csv_reader)

        if not rows:
            return "", {'rows': 0, 'columns': []}

        # Convert CSV to readable text
        content_parts = []
        for i, row in enumerate(rows, 1):
            row_text = f"Row {i}:\n"
            row_text += '\n'.join([f"  {key}: {value}" for key, value in row.items()])
            content_parts.append(row_text)

        content = '\n\n'.join(content_parts)

        metadata = {
            'rows': len(rows),
            'columns': list(rows[0].keys()) if rows else []
        }

        return content, metadata

    def _load_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Load PDF file using pypdf"""
        if PdfReader is None:
            raise ImportError("pypdf is required for PDF support. Install with: pip install pypdf")

        with open(str(file_path), 'rb') as f:
            pdf_reader = PdfReader(f)

            pages = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    pages.append(f"--- Page {page_num} ---\n{text}")

            content = '\n\n'.join(pages)

        metadata = {
            'pages': len(pdf_reader.pages),
            'title': pdf_reader.metadata.title if pdf_reader.metadata else None,
            'author': pdf_reader.metadata.author if pdf_reader.metadata else None
        }

        return content, metadata

    def _load_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Load DOCX file using python-docx"""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")

        doc = DocxDocument(str(file_path))

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        content = '\n\n'.join(paragraphs)

        # Extract tables if any
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(' | '.join(row_data))
            tables_text.append('\n'.join(table_data))

        if tables_text:
            content += '\n\n--- Tables ---\n\n' + '\n\n'.join(tables_text)

        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections)
        }

        return content, metadata

    @staticmethod
    def _dict_to_text(data: Dict, indent: int = 0) -> str:
        """Convert dictionary to readable text"""
        lines = []
        prefix = '  ' * indent

        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(DocumentLoader._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}: {', '.join(map(str, value))}")
            else:
                lines.append(f"{prefix}{key}: {value}")

        return '\n'.join(lines)


# Singleton instance
document_loader = DocumentLoader()